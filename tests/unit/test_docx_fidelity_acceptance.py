import shutil
from pathlib import Path

import pytest
from docx import Document

from docmergeforge.core.exceptions import ValidationError
from docmergeforge.docx import fidelity_acceptance
from docmergeforge.docx.fidelity import FidelityCapability
from docmergeforge.docx.native import NativeCommandResult


def _capability(mode: str) -> FidelityCapability:
    return FidelityCapability(
        mode=mode,
        available=True,
        production_ready=False,
        detail="test adapter",
        automation_ready=True,
        executable="fake-office",
    )


def _write_fixture(path: Path) -> None:
    document = Document()
    document.add_heading("Chapter", level=1)
    document.add_paragraph("Body text")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    section = document.sections[0]
    section.header.paragraphs[0].text = "Header text"
    section.footer.paragraphs[0].text = "Footer text"
    document.save(path)


def test_fidelity_acceptance_records_matching_roundtrip(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _write_fixture(source)

    monkeypatch.setattr(
        fidelity_acceptance,
        "require_fidelity_automation",
        lambda mode: _capability(mode),
    )

    def fake_roundtrip(
        source_path: Path, destination: Path, **kwargs: object
    ) -> NativeCommandResult:
        shutil.copy2(source_path, destination)
        return NativeCommandResult(("fake-office",), "", "")

    monkeypatch.setattr(fidelity_acceptance, "libreoffice_roundtrip_copy", fake_roundtrip)
    evidence = fidelity_acceptance.run_fidelity_roundtrip_acceptance(
        source,
        output,
        "libreoffice",
    )

    assert evidence.accepted
    assert evidence.structure_matches
    assert evidence.content_matches
    assert evidence.source_structure == evidence.output_structure
    assert evidence.source_content == evidence.output_content
    assert evidence.source_structure.header_paragraphs == 1
    assert evidence.source_structure.footer_paragraphs == 1
    assert evidence.to_dict()["accepted"] is True


def test_fidelity_acceptance_flags_structural_change(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _write_fixture(source)

    monkeypatch.setattr(
        fidelity_acceptance,
        "require_fidelity_automation",
        lambda mode: _capability(mode),
    )

    def changed_roundtrip(
        source_path: Path, destination: Path, **kwargs: object
    ) -> NativeCommandResult:
        changed = Document(str(source_path))
        changed.add_paragraph("Unexpected structural change")
        changed.save(destination)
        return NativeCommandResult(("fake-office",), "", "")

    monkeypatch.setattr(fidelity_acceptance, "word_roundtrip_copy", changed_roundtrip)
    evidence = fidelity_acceptance.run_fidelity_roundtrip_acceptance(source, output, "word")

    assert not evidence.accepted
    assert not evidence.structure_matches


def test_fidelity_acceptance_flags_text_change_with_same_structure(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _write_fixture(source)

    monkeypatch.setattr(
        fidelity_acceptance,
        "require_fidelity_automation",
        lambda mode: _capability(mode),
    )

    def changed_roundtrip(
        source_path: Path, destination: Path, **kwargs: object
    ) -> NativeCommandResult:
        changed = Document(str(source_path))
        changed.paragraphs[1].text = "Different body text"
        changed.sections[0].header.paragraphs[0].text = "Different header text"
        changed.save(destination)
        return NativeCommandResult(("fake-office",), "", "")

    monkeypatch.setattr(fidelity_acceptance, "word_roundtrip_copy", changed_roundtrip)
    evidence = fidelity_acceptance.run_fidelity_roundtrip_acceptance(source, output, "word")

    assert evidence.structure_matches
    assert not evidence.content_matches
    assert not evidence.accepted
    assert evidence.source_content != evidence.output_content


def test_fidelity_acceptance_rejects_portable_mode(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_fixture(source)

    with pytest.raises(ValidationError, match="requires mode"):
        fidelity_acceptance.run_fidelity_roundtrip_acceptance(
            source,
            tmp_path / "output.docx",
            "portable",
        )
