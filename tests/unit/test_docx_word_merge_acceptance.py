import shutil
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_SECTION

from docmergeforge.core.exceptions import ValidationError
from docmergeforge.docx import word_merge_acceptance
from docmergeforge.docx.fidelity import FidelityCapability
from docmergeforge.docx.native import NativeCommandResult
from docmergeforge.docx.word_merge import WordNativeMergeResult


def _write_source(path: Path, heading: str, body: str) -> None:
    document = Document()
    document.add_heading(heading, level=1)
    document.add_paragraph(body)
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = f"table-{heading}"
    document.save(path)


def _capability() -> FidelityCapability:
    return FidelityCapability(
        mode="word",
        available=True,
        production_ready=False,
        detail="test Word host",
        automation_ready=True,
        executable="fake-powershell",
    )


def _copy_section_layout(source: object, target: object) -> None:
    for attribute in (
        "start_type",
        "orientation",
        "page_width",
        "page_height",
        "top_margin",
        "bottom_margin",
        "left_margin",
        "right_margin",
        "gutter",
        "header_distance",
        "footer_distance",
        "different_first_page_header_footer",
    ):
        setattr(target, attribute, getattr(source, attribute))
    for name in (
        "header",
        "first_page_header",
        "even_page_header",
        "footer",
        "first_page_footer",
        "even_page_footer",
    ):
        source_story = getattr(source, name)
        target_story = getattr(target, name)
        target_story.is_linked_to_previous = source_story.is_linked_to_previous


def _synthetic_merge(sources: tuple[Path, ...], output: Path) -> None:
    merged = Document()
    merged._body.clear_content()
    for index, source in enumerate(sources):
        current = Document(str(source))
        target_section = (
            merged.sections[0]
            if index == 0
            else merged.add_section(WD_SECTION.NEW_PAGE)
        )
        _copy_section_layout(current.sections[0], target_section)
        for paragraph in current.paragraphs:
            target = merged.add_paragraph(style=paragraph.style.name)
            target.add_run(paragraph.text)
        for table in current.tables:
            target_table = merged.add_table(rows=len(table.rows), cols=len(table.columns))
            for row_index, row in enumerate(table.rows):
                for column_index, cell in enumerate(row.cells):
                    target_table.cell(row_index, column_index).text = cell.text
    merged.save(output)


def test_word_merge_acceptance_accepts_measured_synthetic_merge(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    first = tmp_path / "Part 1.docx"
    second = tmp_path / "Part 2.docx"
    output = tmp_path / "merged.docx"
    _write_source(first, "Part 1", "Alpha")
    _write_source(second, "Part 2", "Beta")

    monkeypatch.setattr(
        word_merge_acceptance,
        "require_fidelity_automation",
        lambda mode: _capability(),
    )

    def fake_merge(
        sources: tuple[Path, ...], destination: Path, **kwargs: object
    ) -> WordNativeMergeResult:
        _synthetic_merge(sources, destination)
        return WordNativeMergeResult(
            source_count=len(sources),
            output=destination,
            command=NativeCommandResult(("fake-powershell",), "", ""),
        )

    monkeypatch.setattr(word_merge_acceptance, "word_merge_documents", fake_merge)
    evidence = word_merge_acceptance.run_word_merge_acceptance(
        [first, second],
        output,
    )

    assert evidence.accepted
    assert evidence.structure_matches
    assert evidence.content_matches
    assert evidence.source_count == 2
    assert evidence.output_sha256
    assert evidence.expected_content.section_properties_sha256
    assert evidence.expected_content.page_number_properties_sha256
    assert evidence.to_dict()["accepted"] is True


def test_word_merge_acceptance_rejects_missing_source_content(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    first = tmp_path / "Part 1.docx"
    second = tmp_path / "Part 2.docx"
    output = tmp_path / "merged.docx"
    _write_source(first, "Part 1", "Alpha")
    _write_source(second, "Part 2", "Beta")

    monkeypatch.setattr(
        word_merge_acceptance,
        "require_fidelity_automation",
        lambda mode: _capability(),
    )

    def incomplete_merge(
        sources: tuple[Path, ...], destination: Path, **kwargs: object
    ) -> WordNativeMergeResult:
        shutil.copy2(sources[0], destination)
        return WordNativeMergeResult(
            source_count=len(sources),
            output=destination,
            command=NativeCommandResult(("fake-powershell",), "", ""),
        )

    monkeypatch.setattr(word_merge_acceptance, "word_merge_documents", incomplete_merge)
    evidence = word_merge_acceptance.run_word_merge_acceptance(
        [first, second],
        output,
    )

    assert not evidence.accepted
    assert not evidence.structure_matches
    assert not evidence.content_matches


def test_word_merge_acceptance_rejects_source_revision_change_before_merge(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = tmp_path / "Part 1.docx"
    output = tmp_path / "merged.docx"
    _write_source(source, "Part 1", "Alpha")
    original_expected_content = word_merge_acceptance.expected_word_merge_content

    def mutating_expected_content(
        sources: tuple[Path, ...],
    ) -> word_merge_acceptance.WordMergeContentSnapshot:
        snapshot = original_expected_content(sources)
        _write_source(sources[0], "Part 1", "Changed")
        return snapshot

    monkeypatch.setattr(
        word_merge_acceptance,
        "expected_word_merge_content",
        mutating_expected_content,
    )

    with pytest.raises(ValidationError, match="Source integrity violation"):
        word_merge_acceptance.run_word_merge_acceptance([source], output)

    assert not output.exists()


def test_word_merge_acceptance_rejects_duplicate_source_before_capability(
    tmp_path: Path,
) -> None:
    source = tmp_path / "Part 1.docx"
    _write_source(source, "Part 1", "Alpha")

    with pytest.raises(ValidationError, match="Duplicate Word merge acceptance source"):
        word_merge_acceptance.run_word_merge_acceptance(
            [source, source],
            tmp_path / "merged.docx",
        )


def test_word_merge_acceptance_rejects_empty_source_set(tmp_path: Path) -> None:
    with pytest.raises(ValidationError, match="at least one DOCX source"):
        word_merge_acceptance.run_word_merge_acceptance([], tmp_path / "merged.docx")
