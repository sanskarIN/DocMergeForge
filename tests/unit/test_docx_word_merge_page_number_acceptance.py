from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from docx.enum.section import WD_SECTION

from docmergeforge.docx import word_merge_acceptance
from docmergeforge.docx.fidelity import FidelityCapability
from docmergeforge.docx.native import NativeCommandResult
from docmergeforge.docx.word_merge import WordNativeMergeResult


def _write_source(path: Path, heading: str, body: str) -> None:
    document = Document()
    document.add_heading(heading, level=1)
    document.add_paragraph(body)
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = f"table-{heading}"
    document.save(path)


def _inject_page_number_properties(path: Path, *, start: str, fmt: str) -> None:
    with ZipFile(path, "r") as source:
        members = {name: source.read(name) for name in source.namelist()}
    document_xml = members["word/document.xml"]
    marker = b"<w:pgSz"
    if marker not in document_xml:
        raise AssertionError("Expected page-size marker in generated DOCX fixture")
    page_number = f'<w:pgNumType w:start="{start}" w:fmt="{fmt}"/>'.encode()
    members["word/document.xml"] = document_xml.replace(
        marker, page_number + marker, 1
    )
    with ZipFile(path, "w", compression=ZIP_DEFLATED) as output:
        for name, payload in members.items():
            output.writestr(name, payload)


def _copy_section_layout(source: object, target: object) -> None:
    for attribute in (
        "orientation",
        "page_width",
        "page_height",
        "top_margin",
        "bottom_margin",
        "left_margin",
        "right_margin",
        "gutter",
        "header_distance",
        "footer_distance",
        "different_first_page_header_footer",
    ):
        setattr(target, attribute, getattr(source, attribute))


def _merge_without_page_number_properties(
    sources: tuple[Path, ...], output: Path
) -> None:
    merged = Document()
    for index, source in enumerate(sources):
        current = Document(str(source))
        target_section = (
            merged.sections[0]
            if index == 0
            else merged.add_section(WD_SECTION.NEW_PAGE)
        )
        _copy_section_layout(current.sections[0], target_section)
        for paragraph in current.paragraphs:
            target = merged.add_paragraph(style=paragraph.style.name)
            target.add_run(paragraph.text)
        for table in current.tables:
            target_table = merged.add_table(rows=len(table.rows), cols=len(table.columns))
            for row_index, row in enumerate(table.rows):
                for column_index, cell in enumerate(row.cells):
                    target_table.cell(row_index, column_index).text = cell.text
    merged.save(output)


def _capability() -> FidelityCapability:
    return FidelityCapability(
        mode="word",
        available=True,
        production_ready=False,
        detail="test Word host",
        automation_ready=True,
        executable="fake-powershell",
    )


def test_word_merge_acceptance_rejects_lost_page_number_restart_and_format(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    first = tmp_path / "Part 1.docx"
    second = tmp_path / "Part 2.docx"
    output = tmp_path / "merged.docx"
    _write_source(first, "Part 1", "Alpha")
    _write_source(second, "Part 2", "Beta")
    _inject_page_number_properties(first, start="1", fmt="decimal")
    _inject_page_number_properties(second, start="7", fmt="upperRoman")

    monkeypatch.setattr(
        word_merge_acceptance,
        "require_fidelity_automation",
        lambda mode: _capability(),
    )

    def fake_merge(
        sources: tuple[Path, ...], destination: Path, **kwargs: object
    ) -> WordNativeMergeResult:
        _merge_without_page_number_properties(sources, destination)
        return WordNativeMergeResult(
            source_count=len(sources),
            output=destination,
            command=NativeCommandResult(("fake-powershell",), "", ""),
        )

    monkeypatch.setattr(word_merge_acceptance, "word_merge_documents", fake_merge)
    evidence = word_merge_acceptance.run_word_merge_acceptance(
        [first, second], output
    )

    assert evidence.structure_matches
    assert not evidence.content_matches
    assert not evidence.accepted
    assert (
        evidence.expected_content.body_paragraphs_sha256
        == evidence.output_content.body_paragraphs_sha256
    )
    assert evidence.expected_content.tables_sha256 == evidence.output_content.tables_sha256
    assert (
        evidence.expected_content.section_properties_sha256
        == evidence.output_content.section_properties_sha256
    )
    assert (
        evidence.expected_content.page_number_properties_sha256
        != evidence.output_content.page_number_properties_sha256
    )
