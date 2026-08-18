import shutil
from pathlib import Path

import pytest
from docx import Document

from docmergeforge.core.exceptions import ValidationError
from docmergeforge.docx import word
from docmergeforge.docx.native import NativeCommandResult


def test_word_roundtrip_copy_uses_separate_validated_output(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = tmp_path / "Part 1.docx"
    destination = tmp_path / "accepted" / "Part 1.docx"
    document = Document()
    document.add_heading("Part 1", level=1)
    document.add_paragraph("Preserve this source.")
    document.save(source)
    before = source.read_bytes()
    captured_script = ""

    def fake_run(command: list[str], **kwargs: object) -> NativeCommandResult:
        nonlocal captured_script
        script_path = Path(command[command.index("-File") + 1])
        captured_script = script_path.read_text(encoding="utf-8")
        destination_index = command.index("-Destination") + 1
        temporary_output = Path(command[destination_index])
        shutil.copy2(source, temporary_output)
        return NativeCommandResult(tuple(command), "saved", "")

    monkeypatch.setattr(word, "run_native_command", fake_run)
    result = word.word_roundtrip_copy(
        source,
        destination,
        powershell="fake-powershell",
    )

    assert destination.exists()
    assert source.read_bytes() == before
    assert result.stdout == "saved"
    assert "$word.AutomationSecurity = 3" in captured_script
    assert "$word.Documents.Open($Source, $false, $true, $false)" in captured_script


def test_word_host_is_not_claimed_off_windows(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(word.sys, "platform", "linux")
    assert word.find_word_powershell_host() is None


def test_word_roundtrip_refuses_same_path(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    Document().save(source)

    with pytest.raises(ValidationError, match="separate output path"):
        word.word_roundtrip_copy(source, source, powershell="fake-powershell")
