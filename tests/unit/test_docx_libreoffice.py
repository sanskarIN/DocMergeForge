import shutil
from pathlib import Path

import pytest
from docx import Document

from docmergeforge.docx import libreoffice
from docmergeforge.docx.native import NativeCommandResult


def test_libreoffice_roundtrip_copy_uses_separate_validated_output(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = tmp_path / "Part 1.docx"
    destination = tmp_path / "accepted" / "Part 1.docx"
    document = Document()
    document.add_heading("Part 1", level=1)
    document.add_paragraph("Preserve this source.")
    document.save(source)
    before = source.read_bytes()
    captured_command: list[str] = []

    def fake_run(command: list[str], **kwargs: object) -> NativeCommandResult:
        captured_command.extend(command)
        outdir = Path(command[command.index("--outdir") + 1])
        shutil.copy2(source, outdir / source.name)
        return NativeCommandResult(tuple(command), "converted", "")

    monkeypatch.setattr(libreoffice, "run_native_command", fake_run)
    result = libreoffice.libreoffice_roundtrip_copy(
        source,
        destination,
        executable="fake-libreoffice",
    )

    assert destination.exists()
    assert source.read_bytes() == before
    assert result.stdout == "converted"
    profile_args = [
        item for item in captured_command if item.startswith("-env:UserInstallation=file:")
    ]
    assert len(profile_args) == 1
    assert "docmergeforge-lo-fidelity-" in profile_args[0]


def test_libreoffice_roundtrip_refuses_overwrite(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    destination = tmp_path / "destination.docx"
    Document().save(source)
    Document().save(destination)

    with pytest.raises(FileExistsError, match="Refusing to overwrite"):
        libreoffice.libreoffice_roundtrip_copy(
            source,
            destination,
            executable="fake-libreoffice",
        )
