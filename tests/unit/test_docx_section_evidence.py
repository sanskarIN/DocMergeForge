from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document

from docmergeforge.core.exceptions import ValidationError
from docmergeforge.docx.section_evidence import (
    page_number_properties_sha256,
    page_number_section_records,
)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _rewrite_document_xml(path: Path, replacements: list[tuple[bytes, bytes]]) -> None:
    with ZipFile(path, "r") as source:
        members = {name: source.read(name) for name in source.namelist()}
    document_xml = members["word/document.xml"]
    for old, new in replacements:
        document_xml = document_xml.replace(old, new)
    members["word/document.xml"] = document_xml
    with ZipFile(path, "w", compression=ZIP_DEFLATED) as output:
        for name, payload in members.items():
            output.writestr(name, payload)


def _write_docx(path: Path, title: str) -> None:
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph("Body")
    document.save(path)


def _set_page_number_properties(
    path: Path,
    *,
    start: str,
    fmt: str,
    chapter_style: str,
    chapter_separator: str,
) -> None:
    marker = b"<w:pgSz"
    page_number = (
        f'<w:pgNumType w:start="{start}" w:fmt="{fmt}" '
        f'w:chapStyle="{chapter_style}" w:chapSep="{chapter_separator}"/>'
    ).encode()
    with ZipFile(path, "r") as archive:
        document_xml = archive.read("word/document.xml")
    if marker not in document_xml:
        raise AssertionError("Expected section page-size marker in generated fixture")
    _rewrite_document_xml(path, [(marker, page_number + marker)])


def test_page_number_section_records_reads_explicit_properties(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_docx(source, "Source")
    _set_page_number_properties(
        source,
        start="7",
        fmt="upperRoman",
        chapter_style="2",
        chapter_separator="hyphen",
    )

    records = page_number_section_records(source)

    assert len(records) == 1
    assert records[0].start == "7"
    assert records[0].format == "upperRoman"
    assert records[0].chapter_style == "2"
    assert records[0].chapter_separator == "hyphen"


def test_page_number_section_records_preserves_empty_explicit_state(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.docx"
    _write_docx(source, "Source")

    records = page_number_section_records(source)

    assert len(records) == 1
    assert records[0].to_dict() == {
        "start": "",
        "format": "",
        "chapter_style": "",
        "chapter_separator": "",
    }


def test_page_number_fingerprint_changes_when_restart_changes(tmp_path: Path) -> None:
    first = tmp_path / "first.docx"
    second = tmp_path / "second.docx"
    _write_docx(first, "First")
    _write_docx(second, "Second")
    _set_page_number_properties(
        first,
        start="1",
        fmt="decimal",
        chapter_style="",
        chapter_separator="",
    )
    _set_page_number_properties(
        second,
        start="5",
        fmt="decimal",
        chapter_style="",
        chapter_separator="",
    )

    original = page_number_properties_sha256([first, second])
    _rewrite_document_xml(
        second,
        [(b'w:start="5"', b'w:start="6"')],
    )
    changed = page_number_properties_sha256([first, second])

    assert original != changed


def test_page_number_fingerprint_binds_document_order(tmp_path: Path) -> None:
    first = tmp_path / "first.docx"
    second = tmp_path / "second.docx"
    _write_docx(first, "First")
    _write_docx(second, "Second")
    _set_page_number_properties(
        first,
        start="1",
        fmt="decimal",
        chapter_style="",
        chapter_separator="",
    )
    _set_page_number_properties(
        second,
        start="9",
        fmt="lowerRoman",
        chapter_style="",
        chapter_separator="",
    )

    assert page_number_properties_sha256([first, second]) != page_number_properties_sha256(
        [second, first]
    )


def test_page_number_evidence_rejects_non_docx(tmp_path: Path) -> None:
    source = tmp_path / "source.txt"
    source.write_text("not a docx", encoding="utf-8")

    with pytest.raises(ValidationError, match="requires a DOCX file"):
        page_number_section_records(source)


def test_page_number_evidence_rejects_missing_document_xml(tmp_path: Path) -> None:
    source = tmp_path / "broken.docx"
    with ZipFile(source, "w", compression=ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types />")

    with pytest.raises(ValidationError, match="missing word/document.xml"):
        page_number_section_records(source)


def test_page_number_fingerprint_requires_sources() -> None:
    with pytest.raises(ValidationError, match="at least one DOCX source"):
        page_number_properties_sha256([])
