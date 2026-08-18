import json
from pathlib import Path

import pytest
from docx import Document

from docmergeforge.core.exceptions import ValidationError
from docmergeforge.docx import word_merge
from docmergeforge.docx.native import NativeCommandResult
from docmergeforge.docx.word_process import WordProcessCleanupResult


def _write_docx(path: Path, text: str) -> None:
    document = Document()
    document.add_paragraph(text)
    document.save(path)


def _clean_process_result() -> WordProcessCleanupResult:
    return WordProcessCleanupResult(
        identity_present=True,
        process_found=False,
        terminated=False,
    )


def _write_identity_from_command(command: list[str]) -> Path:
    identity = Path(command[command.index("-ProcessIdentityFile") + 1])
    identity.write_text(
        json.dumps(
            {
                "process_id": 4242,
                "process_name": "WINWORD",
                "start_time_utc_ticks": 638910000000000000,
            }
        ),
        encoding="utf-8",
    )
    return identity


def test_word_native_merge_builds_ordered_manifest_and_preserves_sources(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    first = tmp_path / "Part 1.docx"
    second = tmp_path / "Part 2.docx"
    output = tmp_path / "merged.docx"
    _write_docx(first, "First")
    _write_docx(second, "Second")
    first_before = first.read_bytes()
    second_before = second.read_bytes()
    captured: dict[str, object] = {}

    def fake_run(command: list[str], *, timeout_seconds: int) -> NativeCommandResult:
        captured["command"] = tuple(command)
        captured["timeout"] = timeout_seconds
        manifest = Path(command[command.index("-Manifest") + 1])
        destination = Path(command[command.index("-Destination") + 1])
        sources = json.loads(manifest.read_text(encoding="utf-8"))
        captured["sources"] = sources
        captured["script"] = Path(command[command.index("-File") + 1]).read_text(
            encoding="utf-8"
        )
        _write_identity_from_command(command)

        merged = Document()
        for source in sources:
            current = Document(source)
            for paragraph in current.paragraphs:
                merged.add_paragraph(paragraph.text)
        merged.save(destination)
        return NativeCommandResult(tuple(command), "", "")

    cleanup_calls: list[tuple[Path, str]] = []

    def fake_cleanup(
        identity_file: Path, *, powershell: str
    ) -> WordProcessCleanupResult:
        cleanup_calls.append((identity_file, powershell))
        return _clean_process_result()

    monkeypatch.setattr(word_merge, "run_native_command", fake_run)
    monkeypatch.setattr(word_merge, "cleanup_word_process_identity", fake_cleanup)

    result = word_merge.word_merge_documents(
        [first, second],
        output,
        powershell="fake-powershell",
        timeout_seconds=45,
    )

    assert result.output == output
    assert result.source_count == 2
    assert captured["sources"] == [str(first.resolve()), str(second.resolve())]
    assert captured["timeout"] == 45
    command = captured["command"]
    assert isinstance(command, tuple)
    assert "-ProcessIdentityFile" in command
    script = str(captured["script"])
    assert "GetWindowThreadProcessId" in script
    assert "start_time_utc_ticks" in script
    assert "$range.InsertBreak(2)" in script
    assert "$range.InsertBreak(3)" in script
    assert "$range.InsertFile" in script
    assert cleanup_calls
    assert cleanup_calls[0][1] == "fake-powershell"
    assert first.read_bytes() == first_before
    assert second.read_bytes() == second_before
    assert output.exists()


def test_word_native_merge_uses_continuous_section_break_when_requested(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "merged.docx"
    _write_docx(source, "Body")
    captured_command: list[str] = []

    def fake_run(command: list[str], *, timeout_seconds: int) -> NativeCommandResult:
        captured_command.extend(command)
        destination = Path(command[command.index("-Destination") + 1])
        _write_identity_from_command(command)
        _write_docx(destination, "Body")
        return NativeCommandResult(tuple(command), "", "")

    monkeypatch.setattr(word_merge, "run_native_command", fake_run)
    monkeypatch.setattr(
        word_merge,
        "cleanup_word_process_identity",
        lambda *args, **kwargs: _clean_process_result(),
    )

    word_merge.word_merge_documents(
        [source],
        output,
        powershell="fake-powershell",
        start_each_on_new_page=False,
    )

    value = captured_command[captured_command.index("-StartEachOnNewPage") + 1]
    assert value == "0"


def test_word_native_merge_cleans_exact_process_after_command_failure(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "merged.docx"
    _write_docx(source, "Body")
    cleanup_calls: list[Path] = []

    def failing_run(command: list[str], *, timeout_seconds: int) -> NativeCommandResult:
        _write_identity_from_command(command)
        raise TimeoutError("simulated PowerShell timeout")

    def fake_cleanup(
        identity_file: Path, *, powershell: str
    ) -> WordProcessCleanupResult:
        cleanup_calls.append(identity_file)
        return WordProcessCleanupResult(
            identity_present=True,
            process_found=True,
            terminated=True,
        )

    monkeypatch.setattr(word_merge, "run_native_command", failing_run)
    monkeypatch.setattr(word_merge, "cleanup_word_process_identity", fake_cleanup)

    with pytest.raises(TimeoutError, match="simulated PowerShell timeout"):
        word_merge.word_merge_documents(
            [source],
            output,
            powershell="fake-powershell",
            timeout_seconds=1,
        )

    assert len(cleanup_calls) == 1
    assert not output.exists()


def test_word_native_merge_fails_if_success_required_forced_word_cleanup(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "merged.docx"
    _write_docx(source, "Body")

    def fake_run(command: list[str], *, timeout_seconds: int) -> NativeCommandResult:
        destination = Path(command[command.index("-Destination") + 1])
        _write_identity_from_command(command)
        _write_docx(destination, "Body")
        return NativeCommandResult(tuple(command), "", "")

    monkeypatch.setattr(word_merge, "run_native_command", fake_run)
    monkeypatch.setattr(
        word_merge,
        "cleanup_word_process_identity",
        lambda *args, **kwargs: WordProcessCleanupResult(
            identity_present=True,
            process_found=True,
            terminated=True,
        ),
    )

    with pytest.raises(ValidationError, match="forcibly terminated"):
        word_merge.word_merge_documents(
            [source],
            output,
            powershell="fake-powershell",
        )

    assert not output.exists()


def test_word_native_merge_rejects_missing_process_identity_after_success(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "merged.docx"
    _write_docx(source, "Body")

    def fake_run(command: list[str], *, timeout_seconds: int) -> NativeCommandResult:
        destination = Path(command[command.index("-Destination") + 1])
        _write_docx(destination, "Body")
        return NativeCommandResult(tuple(command), "", "")

    monkeypatch.setattr(word_merge, "run_native_command", fake_run)

    with pytest.raises(ValidationError, match="did not record its Word process identity"):
        word_merge.word_merge_documents(
            [source],
            output,
            powershell="fake-powershell",
        )

    assert not output.exists()


def test_word_native_merge_rejects_duplicate_sources(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_docx(source, "Body")

    with pytest.raises(ValidationError, match="Duplicate Microsoft Word merge source"):
        word_merge.word_merge_documents(
            [source, source],
            tmp_path / "merged.docx",
            powershell="fake-powershell",
        )


def test_word_native_merge_refuses_existing_output(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "merged.docx"
    _write_docx(source, "Body")
    _write_docx(output, "Existing")

    with pytest.raises(FileExistsError, match="Refusing to overwrite"):
        word_merge.word_merge_documents(
            [source],
            output,
            powershell="fake-powershell",
        )


def test_word_native_merge_rejects_source_output_collision(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_docx(source, "Body")

    with pytest.raises(ValidationError, match="separate output path"):
        word_merge.word_merge_documents(
            [source],
            source,
            powershell="fake-powershell",
        )


def test_word_native_merge_rejects_nonpositive_timeout(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_docx(source, "Body")

    with pytest.raises(ValidationError, match="timeout must be at least one second"):
        word_merge.word_merge_documents(
            [source],
            tmp_path / "merged.docx",
            powershell="fake-powershell",
            timeout_seconds=0,
        )
