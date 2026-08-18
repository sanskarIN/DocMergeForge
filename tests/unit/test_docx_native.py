from pathlib import Path
from subprocess import CompletedProcess, TimeoutExpired

import pytest
from docx import Document

from docmergeforge.core.exceptions import ValidationError
from docmergeforge.docx import native
from docmergeforge.utilities.hashing import sha256_file


def test_run_native_command_rejects_empty_command() -> None:
    with pytest.raises(ValidationError, match="cannot be empty"):
        native.run_native_command([])


def test_run_native_command_captures_success(monkeypatch: pytest.MonkeyPatch) -> None:
    def fake_run(*args: object, **kwargs: object) -> CompletedProcess[str]:
        return CompletedProcess(["office"], 0, stdout="ok", stderr="")

    monkeypatch.setattr(native.subprocess, "run", fake_run)
    result = native.run_native_command(["office", "--headless"])
    assert result.command == ("office", "--headless")
    assert result.stdout == "ok"


def test_run_native_command_fails_closed_on_nonzero(monkeypatch: pytest.MonkeyPatch) -> None:
    def fake_run(*args: object, **kwargs: object) -> CompletedProcess[str]:
        return CompletedProcess(["office"], 9, stdout="", stderr="failure")

    monkeypatch.setattr(native.subprocess, "run", fake_run)
    with pytest.raises(ValidationError, match="exit code 9"):
        native.run_native_command(["office"])


def test_run_native_command_fails_closed_on_timeout(monkeypatch: pytest.MonkeyPatch) -> None:
    def fake_run(*args: object, **kwargs: object) -> CompletedProcess[str]:
        raise TimeoutExpired(cmd=["office"], timeout=1)

    monkeypatch.setattr(native.subprocess, "run", fake_run)
    with pytest.raises(ValidationError, match="timed out"):
        native.run_native_command(["office"], timeout_seconds=1)


def test_validate_native_docx_output_and_source_hash(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("original")
    document.save(source)
    expected = sha256_file(source)

    native.validate_native_docx_output(source)
    native.verify_native_source_unchanged(source, expected)

    changed = Document(str(source))
    changed.add_paragraph("changed")
    changed.save(source)
    with pytest.raises(ValidationError, match="Source integrity violation"):
        native.verify_native_source_unchanged(source, expected)


def test_promote_validated_native_output_succeeds(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    temporary = tmp_path / "temporary.docx"
    destination = tmp_path / "final.docx"
    Document().save(source)
    result = Document()
    result.add_paragraph("result")
    result.save(temporary)

    native.promote_validated_native_docx_output(
        temporary,
        destination,
        {source: sha256_file(source)},
    )

    assert destination.exists()
    assert not temporary.exists()


def test_promote_validated_native_output_removes_destination_on_final_failure(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "source.docx"
    temporary = tmp_path / "temporary.docx"
    destination = tmp_path / "final.docx"
    Document().save(source)
    Document().save(temporary)
    checks = 0

    original_verify = native.verify_native_sources_unchanged

    def fail_second_check(source_hashes: dict[Path, str]) -> None:
        nonlocal checks
        checks += 1
        original_verify(source_hashes)
        if checks == 2:
            raise ValidationError("Source integrity violation after promotion")

    monkeypatch.setattr(native, "verify_native_sources_unchanged", fail_second_check)

    with pytest.raises(ValidationError, match="after promotion"):
        native.promote_validated_native_docx_output(
            temporary,
            destination,
            {source: sha256_file(source)},
        )

    assert checks == 2
    assert not destination.exists()
