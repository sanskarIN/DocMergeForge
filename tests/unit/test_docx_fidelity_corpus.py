from pathlib import Path

import pytest
from docx import Document

from docmergeforge.core.exceptions import ValidationError
from docmergeforge.docx import fidelity_corpus
from docmergeforge.docx.fidelity_acceptance import (
    FidelityAcceptanceEvidence,
    snapshot_docx_content,
    snapshot_docx_structure,
)
from docmergeforge.utilities.hashing import sha256_file


def _write_docx(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(text, level=1)
    document.add_paragraph("body")
    document.save(path)


def _accepted_evidence(source: Path, output: Path, mode: str) -> FidelityAcceptanceEvidence:
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_bytes(source.read_bytes())
    structure = snapshot_docx_structure(source)
    content = snapshot_docx_content(source)
    return FidelityAcceptanceEvidence(
        mode=mode,
        source=source,
        output=output,
        source_sha256=sha256_file(source),
        output_sha256=sha256_file(output),
        source_structure=structure,
        output_structure=structure,
        source_content=content,
        output_content=content,
        source_risks=(),
        output_risks=(),
        structure_matches=True,
        content_matches=True,
        new_risks=(),
    )


def test_discover_fidelity_corpus_is_recursive_and_deterministic(tmp_path: Path) -> None:
    corpus = tmp_path / "corpus"
    _write_docx(corpus / "b.docx", "B")
    _write_docx(corpus / "nested" / "a.docx", "A")
    (corpus / "ignore.txt").write_text("ignore", encoding="utf-8")

    discovered = fidelity_corpus.discover_fidelity_corpus(corpus)

    assert [path.relative_to(corpus).as_posix() for path in discovered] == [
        "b.docx",
        "nested/a.docx",
    ]


def test_run_fidelity_corpus_keeps_report_paths_relative(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    corpus = tmp_path / "private-corpus"
    output = tmp_path / "evidence"
    _write_docx(corpus / "nested" / "sample.docx", "Sample")

    monkeypatch.setattr(
        fidelity_corpus,
        "run_fidelity_roundtrip_acceptance",
        lambda source, destination, mode, **kwargs: _accepted_evidence(
            source, destination, mode
        ),
    )
    report = fidelity_corpus.run_fidelity_corpus(corpus, output, "libreoffice")
    payload = report.to_dict()

    assert report.accepted
    assert payload["discovered_count"] == 1
    assert payload["processed_count"] == 1
    item = payload["items"][0]
    assert item["source"] == "nested/sample.docx"
    assert item["output"] == "roundtrip/nested/sample.docx"
    assert item["evidence"]["content_matches"] is True
    assert str(corpus.resolve()) not in str(payload)
    assert str(output.resolve()) not in str(payload)


def test_corpus_errors_redact_absolute_private_roots(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    corpus = tmp_path / "private-corpus"
    output = tmp_path / "private-evidence"
    source = corpus / "nested" / "sample.docx"
    _write_docx(source, "Sample")

    def fail_acceptance(
        source_path: Path, destination: Path, *args: object, **kwargs: object
    ) -> FidelityAcceptanceEvidence:
        raise ValidationError(f"failed {source_path} -> {destination}")

    monkeypatch.setattr(
        fidelity_corpus,
        "run_fidelity_roundtrip_acceptance",
        fail_acceptance,
    )
    report = fidelity_corpus.run_fidelity_corpus(corpus, output, "libreoffice")
    payload = report.to_dict()
    error = payload["items"][0]["error"]

    assert isinstance(error, str)
    assert str(corpus.resolve()) not in error
    assert str(output.resolve()) not in error
    assert "<corpus>/nested/sample.docx" in error
    assert "<evidence>/roundtrip/nested/sample.docx" in error


def test_fail_fast_partial_corpus_is_never_accepted(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    corpus = tmp_path / "corpus"
    output = tmp_path / "evidence"
    _write_docx(corpus / "a.docx", "A")
    _write_docx(corpus / "b.docx", "B")

    def fail_acceptance(*args: object, **kwargs: object) -> FidelityAcceptanceEvidence:
        raise ValidationError("adapter failed")

    monkeypatch.setattr(
        fidelity_corpus,
        "run_fidelity_roundtrip_acceptance",
        fail_acceptance,
    )
    report = fidelity_corpus.run_fidelity_corpus(
        corpus,
        output,
        "libreoffice",
        fail_fast=True,
    )

    assert report.discovered_count == 2
    assert report.processed_count == 1
    assert report.stopped_early
    assert not report.accepted


def test_run_fidelity_corpus_refuses_output_inside_source(tmp_path: Path) -> None:
    corpus = tmp_path / "corpus"
    _write_docx(corpus / "sample.docx", "Sample")

    with pytest.raises(ValidationError, match="outside the source corpus"):
        fidelity_corpus.run_fidelity_corpus(
            corpus,
            corpus / "evidence",
            "libreoffice",
        )


def test_write_fidelity_corpus_report_refuses_overwrite(tmp_path: Path) -> None:
    report = fidelity_corpus.FidelityCorpusReport(
        mode="libreoffice",
        pattern="*.docx",
        recursive=True,
        discovered_count=0,
        items=(),
    )
    destination = tmp_path / "report.json"
    destination.write_text("existing", encoding="utf-8")

    with pytest.raises(FileExistsError, match="Refusing to overwrite"):
        fidelity_corpus.write_fidelity_corpus_report(report, destination)
