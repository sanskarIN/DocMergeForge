import shutil
from pathlib import Path

import pytest
from docx import Document

from docmergeforge.core.exceptions import ValidationError
from docmergeforge.docx import libreoffice_merge_acceptance
from docmergeforge.docx.fidelity import FidelityCapability
from docmergeforge.docx.libreoffice_merge import LibreOfficeNativeMergeResult


def _write_source(path: Path, title: str, body: str) -> None:
    document = Document()
    document.add_paragraph(title)
    document.add_paragraph(body)
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = f"table-{title}"
    document.save(path)


def _capability() -> FidelityCapability:
    return FidelityCapability(
        mode="libreoffice",
        available=True,
        production_ready=False,
        detail="test LibreOffice host",
        automation_ready=True,
        executable="fake-soffice",
    )


def _synthetic_merge(sources: tuple[Path, ...], output: Path) -> None:
    merged = Document()
    merged._body.clear_content()
    for source in sources:
        current = Document(str(source))
        for paragraph in current.paragraphs:
            merged.add_paragraph(paragraph.text)
        for table in current.tables:
            target = merged.add_table(rows=len(table.rows), cols=len(table.columns))
            for row_index, row in enumerate(table.rows):
                for column_index, cell in enumerate(row.cells):
                    target.cell(row_index, column_index).text = cell.text
    merged.save(output)


def test_libreoffice_merge_acceptance_accepts_measured_synthetic_merge(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    first = tmp_path / "Part 1.docx"
    second = tmp_path / "Part 2.docx"
    output = tmp_path / "merged.docx"
    _write_source(first, "Part 1", "Alpha")
    _write_source(second, "Part 2", "Beta")

    monkeypatch.setattr(
        libreoffice_merge_acceptance,
        "require_fidelity_automation",
        lambda mode: _capability(),
    )
    monkeypatch.setattr(
        libreoffice_merge_acceptance,
        "find_uno_python",
        lambda: "fake-python",
    )

    def fake_merge(
        sources: tuple[Path, ...], destination: Path, **kwargs: object
    ) -> LibreOfficeNativeMergeResult:
        _synthetic_merge(sources, destination)
        return LibreOfficeNativeMergeResult(
            source_count=len(sources),
            output=destination,
            worker_stdout="",
            worker_stderr="",
        )

    monkeypatch.setattr(
        libreoffice_merge_acceptance,
        "libreoffice_merge_documents",
        fake_merge,
    )
    evidence = libreoffice_merge_acceptance.run_libreoffice_merge_acceptance(
        [first, second],
        output,
    )

    assert evidence.accepted
    assert evidence.structure_matches
    assert evidence.content_matches
    assert evidence.source_count == 2
    assert evidence.output_sha256
    assert evidence.to_dict()["accepted"] is True


def test_libreoffice_merge_acceptance_rejects_missing_source_content(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    first = tmp_path / "Part 1.docx"
    second = tmp_path / "Part 2.docx"
    output = tmp_path / "merged.docx"
    _write_source(first, "Part 1", "Alpha")
    _write_source(second, "Part 2", "Beta")

    monkeypatch.setattr(
        libreoffice_merge_acceptance,
        "require_fidelity_automation",
        lambda mode: _capability(),
    )
    monkeypatch.setattr(
        libreoffice_merge_acceptance,
        "find_uno_python",
        lambda: "fake-python",
    )

    def incomplete_merge(
        sources: tuple[Path, ...], destination: Path, **kwargs: object
    ) -> LibreOfficeNativeMergeResult:
        shutil.copy2(sources[0], destination)
        return LibreOfficeNativeMergeResult(
            source_count=len(sources),
            output=destination,
            worker_stdout="",
            worker_stderr="",
        )

    monkeypatch.setattr(
        libreoffice_merge_acceptance,
        "libreoffice_merge_documents",
        incomplete_merge,
    )
    evidence = libreoffice_merge_acceptance.run_libreoffice_merge_acceptance(
        [first, second],
        output,
    )

    assert not evidence.accepted
    assert not evidence.structure_matches
    assert not evidence.content_matches


def test_libreoffice_merge_acceptance_rejects_source_revision_change(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "Part 1.docx"
    output = tmp_path / "merged.docx"
    _write_source(source, "Part 1", "Alpha")
    original_expected = libreoffice_merge_acceptance.expected_libreoffice_merge_content

    def mutating_expected(
        sources: tuple[Path, ...],
    ) -> libreoffice_merge_acceptance.LibreOfficeMergeContentSnapshot:
        snapshot = original_expected(sources)
        _write_source(sources[0], "Part 1", "Changed")
        return snapshot

    monkeypatch.setattr(
        libreoffice_merge_acceptance,
        "expected_libreoffice_merge_content",
        mutating_expected,
    )

    with pytest.raises(ValidationError, match="Source integrity violation"):
        libreoffice_merge_acceptance.run_libreoffice_merge_acceptance(
            [source],
            output,
        )

    assert not output.exists()


def test_libreoffice_merge_acceptance_requires_uno_python(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "Part 1.docx"
    _write_source(source, "Part 1", "Alpha")
    monkeypatch.setattr(
        libreoffice_merge_acceptance,
        "require_fidelity_automation",
        lambda mode: _capability(),
    )
    monkeypatch.setattr(
        libreoffice_merge_acceptance,
        "find_uno_python",
        lambda: None,
    )

    with pytest.raises(ValidationError, match="Python UNO bridge"):
        libreoffice_merge_acceptance.run_libreoffice_merge_acceptance(
            [source],
            tmp_path / "merged.docx",
        )


def test_libreoffice_merge_acceptance_rejects_duplicate_sources(tmp_path: Path) -> None:
    source = tmp_path / "Part 1.docx"
    _write_source(source, "Part 1", "Alpha")

    with pytest.raises(ValidationError, match="Duplicate LibreOffice acceptance source"):
        libreoffice_merge_acceptance.run_libreoffice_merge_acceptance(
            [source, source],
            tmp_path / "merged.docx",
        )
