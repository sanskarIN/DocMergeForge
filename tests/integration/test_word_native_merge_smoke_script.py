import json
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

from docmergeforge.docx.section_evidence import page_number_section_records
from docmergeforge.docx.word_merge_acceptance import (
    WordMergeAcceptanceEvidence,
    WordMergeContentSnapshot,
    WordMergeStructureSnapshot,
)
from scripts import check_word_native_merge_smoke as script


def _accepted_evidence(output: Path) -> WordMergeAcceptanceEvidence:
    structure = WordMergeStructureSnapshot(
        paragraphs=8,
        tables=2,
        inline_shapes=0,
        headings=2,
        sections=2,
        header_paragraphs=2,
        footer_paragraphs=2,
        header_tables=0,
        footer_tables=0,
    )
    content = WordMergeContentSnapshot(
        body_paragraphs_sha256="a" * 64,
        tables_sha256="b" * 64,
        headers_sha256="c" * 64,
        footers_sha256="d" * 64,
        section_properties_sha256="1" * 64,
        page_number_properties_sha256="2" * 64,
    )
    return WordMergeAcceptanceEvidence(
        source_count=2,
        source_sha256=("e" * 64, "f" * 64),
        output=output,
        output_sha256="0" * 64,
        expected_structure=structure,
        output_structure=structure,
        expected_content=content,
        output_content=content,
        source_risks=(),
        output_risks=(),
        new_risks=(),
        structure_matches=True,
        content_matches=True,
    )


def test_word_native_merge_smoke_builds_two_distinct_sources(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    output_dir = tmp_path / "evidence"
    captured_sources: list[Path] = []

    def fake_acceptance(
        sources: list[Path], destination: Path, **kwargs: object
    ) -> WordMergeAcceptanceEvidence:
        captured_sources.extend(sources)
        return _accepted_evidence(destination)

    monkeypatch.setattr(script, "run_word_merge_acceptance", fake_acceptance)
    exit_code = script.main(["--output-dir", str(output_dir), "--timeout", "30"])

    first, second = captured_sources
    first_doc = Document(str(first))
    second_doc = Document(str(second))
    second_section = second_doc.sections[0]
    first_page_number = page_number_section_records(first)[0]
    second_page_number = page_number_section_records(second)[0]
    payload = json.loads(
        (output_dir / "word-native-merge-evidence.json").read_text(encoding="utf-8")
    )

    assert exit_code == 0
    assert first_doc.sections[0].orientation != WD_ORIENT.LANDSCAPE
    assert second_section.orientation == WD_ORIENT.LANDSCAPE
    assert second_section.left_margin == Inches(0.70)
    assert second_section.right_margin == Inches(0.80)
    assert second_section.top_margin == Inches(0.60)
    assert second_section.bottom_margin == Inches(0.65)
    assert second_section.header_distance == Inches(0.25)
    assert second_section.footer_distance == Inches(0.30)
    assert first_doc.sections[0].header.paragraphs[0].text.endswith("Source 1")
    assert second_section.header.paragraphs[0].text.endswith("Source 2")
    assert first_page_number.start == "1"
    assert first_page_number.format == "decimal"
    assert second_page_number.start == "7"
    assert second_page_number.format == "upperRoman"
    assert payload["expected_content"]["section_properties_sha256"] == "1" * 64
    assert payload["expected_content"]["page_number_properties_sha256"] == "2" * 64
    assert payload["accepted"] is True


def test_word_native_merge_smoke_refuses_artifact_overwrite(tmp_path: Path) -> None:
    output_dir = tmp_path / "evidence"
    output_dir.mkdir()
    (output_dir / "word-native-merged.docx").write_bytes(b"existing")

    with pytest.raises(SystemExit, match="Refusing to overwrite existing Word smoke artifact"):
        script.main(["--output-dir", str(output_dir)])
