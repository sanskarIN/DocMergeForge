from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfReader, PdfWriter

from docmergeforge.core.models import DocumentKind
from docmergeforge.web.app import create_app, output_filename, safe_filename


def _pdf_bytes() -> bytes:
    buffer = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.write(buffer)
    return buffer.getvalue()


def _docx_bytes(text: str) -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph(text)
    document.save(buffer)
    return buffer.getvalue()


def test_safe_filename_drops_path_components() -> None:
    assert safe_filename("../../Part 1.pdf") == "Part 1.pdf"
    assert safe_filename(r"..\..\Part 1.pdf") == "Part 1.pdf"
    assert output_filename("../unsafe", kind=DocumentKind.PDF) == "unsafe.pdf"


def test_browser_shell_uses_fragment_token_and_security_headers() -> None:
    client = TestClient(create_app())
    response = client.get("/")

    assert response.status_code == 200
    assert 'id="access-token"' in response.text
    assert 'fragment.get("token")' in response.text
    assert "queryToken" not in response.text
    assert response.headers["referrer-policy"] == "no-referrer"
    assert response.headers["x-content-type-options"] == "nosniff"
    assert response.headers["x-frame-options"] == "DENY"
    assert "frame-ancestors 'none'" in response.headers["content-security-policy"]


def test_health_and_platform_routes() -> None:
    client = TestClient(create_app())
    assert client.get("/healthz").json()["status"] == "ok"
    ids = {item["id"] for item in client.get("/api/platforms").json()["targets"]}
    assert {"windows", "macos", "linux", "android", "ios", "web"} <= ids


def test_merge_pdf_from_browser_uploads() -> None:
    client = TestClient(create_app())
    payload = _pdf_bytes()
    response = client.post(
        "/api/merge",
        files=[
            ("files", ("Part 10.pdf", payload, "application/pdf")),
            ("files", ("Part 2.pdf", payload, "application/pdf")),
        ],
        data={"output_name": "Book Master"},
    )
    assert response.status_code == 200, response.text
    assert response.headers["content-type"].startswith("application/pdf")
    assert len(PdfReader(BytesIO(response.content)).pages) == 2


def test_merge_docx_from_browser_uploads() -> None:
    client = TestClient(create_app())
    response = client.post(
        "/api/merge",
        files=[
            (
                "files",
                (
                    "Part 1.docx",
                    _docx_bytes("First"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            ),
            (
                "files",
                (
                    "Part 2.docx",
                    _docx_bytes("Second"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            ),
        ],
    )
    assert response.status_code == 200, response.text
    merged = Document(BytesIO(response.content))
    text = "\n".join(paragraph.text for paragraph in merged.paragraphs)
    assert "First" in text
    assert "Second" in text


def test_access_token_is_required_when_configured() -> None:
    client = TestClient(create_app(access_token="secret-token"))
    payload = _pdf_bytes()
    denied = client.post(
        "/api/merge",
        files=[("files", ("Part 1.pdf", payload, "application/pdf"))],
    )
    assert denied.status_code == 401

    allowed = client.post(
        "/api/merge",
        files=[("files", ("Part 1.pdf", payload, "application/pdf"))],
        headers={"X-DocMergeForge-Token": "secret-token"},
    )
    assert allowed.status_code == 200


def test_merge_failure_does_not_expose_internal_exception(monkeypatch) -> None:
    def fail_merge(*args: object, **kwargs: object) -> None:
        raise RuntimeError("sensitive-host-path:/private/manuscript")

    monkeypatch.setattr("docmergeforge.web.app.PdfMergeEngine.merge", fail_merge)
    client = TestClient(create_app())
    response = client.post(
        "/api/merge",
        files=[("files", ("Part 1.pdf", _pdf_bytes(), "application/pdf"))],
    )

    assert response.status_code == 422
    assert response.json()["detail"] == (
        "Merge failed. Check the DocMergeForge host logs for details."
    )
    assert "sensitive-host-path" not in response.text


def test_upload_size_limit_is_fail_closed() -> None:
    client = TestClient(create_app(max_upload_bytes=1))
    response = client.post(
        "/api/merge",
        files=[("files", ("Part 1.pdf", _pdf_bytes(), "application/pdf"))],
    )
    assert response.status_code == 413


def test_mixed_formats_are_rejected() -> None:
    client = TestClient(create_app())
    response = client.post(
        "/api/merge",
        files=[
            ("files", ("Part 1.pdf", _pdf_bytes(), "application/pdf")),
            (
                "files",
                (
                    "Part 2.docx",
                    _docx_bytes("Two"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            ),
        ],
    )
    assert response.status_code == 400
