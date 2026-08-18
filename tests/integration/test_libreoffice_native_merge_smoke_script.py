import json
from pathlib import Path

import pytest
from docx import Document

from docmergeforge.docx.libreoffice_merge_acceptance import (
    LibreOfficeMergeAcceptanceEvidence,
    LibreOfficeMergeContentSnapshot,
    LibreOfficeMergeStructureSnapshot,
)
from scripts import check_libreoffice_native_merge_smoke as script


def _accepted_evidence(output: Path) -> LibreOfficeMergeAcceptanceEvidence:
    structure = LibreOfficeMergeStructureSnapshot(
        paragraphs=4,
        tables=2,
        inline_shapes=0,
        headings=0,
    )
    content = LibreOfficeMergeContentSnapshot(
        body_paragraphs_sha256="a" * 64,
        tables_sha256="b" * 64,
    )
    return LibreOfficeMergeAcceptanceEvidence(
        source_count=2,
        source_sha256=("c" * 64, "d" * 64),
        output=output,
        output_sha256="e" * 64,
        expected_structure=structure,
        output_structure=structure,
        expected_content=content,
        output_content=content,
        source_risks=(),
        output_risks=(),
        new_risks=(),
        structure_matches=True,
        content_matches=True,
    )


def test_libreoffice_native_merge_smoke_builds_two_distinct_sources(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    output_dir = tmp_path / "evidence"
    captured_sources: list[Path] = []

    def fake_acceptance(
        sources: list[Path], destination: Path, **kwargs: object
    ) -> LibreOfficeMergeAcceptanceEvidence:
        captured_sources.extend(sources)
        return _accepted_evidence(destination)

    monkeypatch.setattr(script, "run_libreoffice_merge_acceptance", fake_acceptance)
    exit_code = script.main(["--output-dir", str(output_dir), "--timeout", "30"])

    first, second = captured_sources
    first_doc = Document(str(first))
    second_doc = Document(str(second))
    payload = json.loads(
        (output_dir / "libreoffice-native-merge-evidence.json").read_text(
            encoding="utf-8"
        )
    )

    assert exit_code == 0
    assert first_doc.paragraphs[0].text.endswith("Source 1")
    assert second_doc.paragraphs[0].text.endswith("Source 2")
    assert first_doc.tables[0].cell(1, 1).text.endswith("Source 1")
    assert second_doc.tables[0].cell(1, 1).text.endswith("Source 2")
    assert payload["accepted"] is True


def test_libreoffice_native_merge_smoke_refuses_artifact_overwrite(
    tmp_path: Path,
) -> None:
    output_dir = tmp_path / "evidence"
    output_dir.mkdir()
    (output_dir / "libreoffice-native-merged.docx").write_bytes(b"existing")

    with pytest.raises(
        SystemExit,
        match="Refusing to overwrite existing LibreOffice smoke artifact",
    ):
        script.main(["--output-dir", str(output_dir)])
