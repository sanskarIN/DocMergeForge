from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document

from docmergeforge.docx.libreoffice_uno_acceptance import (
    run_libreoffice_uno_acceptance,
)


def _write_fixture(path: Path, title: str) -> None:
    document = Document()
    document.core_properties.title = title
    document.add_paragraph(title)
    document.add_paragraph(f"Representative supervised UNO body for {title}.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Feature"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Source"
    table.cell(1, 1).text = title
    document.save(path)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Run the supervised LibreOffice UNO multi-DOCX merge smoke."
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("libreoffice-uno-evidence"),
    )
    parser.add_argument("--timeout", type=int, default=300)
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    if args.timeout < 1:
        raise SystemExit("--timeout must be at least one second")

    output_dir: Path = args.output_dir
    output_dir.mkdir(parents=True, exist_ok=True)
    first = output_dir / "libreoffice-uno-source-01.docx"
    second = output_dir / "libreoffice-uno-source-02.docx"
    merged = output_dir / "libreoffice-uno-merged.docx"
    evidence_path = output_dir / "libreoffice-uno-merge-evidence.json"
    for path in (first, second, merged, evidence_path):
        if path.exists():
            raise SystemExit(
                f"Refusing to overwrite existing LibreOffice UNO smoke artifact: {path}"
            )

    _write_fixture(first, "LibreOffice UNO Source 1")
    _write_fixture(second, "LibreOffice UNO Source 2")
    evidence = run_libreoffice_uno_acceptance(
        [first, second],
        merged,
        timeout_seconds=args.timeout,
        start_each_on_new_page=True,
    )
    evidence_path.write_text(
        json.dumps(evidence.to_dict(), indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    print(evidence_path)
    return 0 if evidence.accepted else 2


if __name__ == "__main__":
    raise SystemExit(main())
