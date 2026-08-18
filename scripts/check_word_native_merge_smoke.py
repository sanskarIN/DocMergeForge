from __future__ import annotations

import argparse
import json
import tempfile
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

from docmergeforge.docx.word_merge_acceptance import run_word_merge_acceptance


def _inject_page_number_properties(path: Path, *, start: int, fmt: str) -> None:
    with ZipFile(path, "r") as source:
        members = {name: source.read(name) for name in source.namelist()}

    document_xml = members["word/document.xml"]
    marker = b"<w:cols"
    if marker not in document_xml:
        raise RuntimeError("Generated Word smoke fixture has no section columns marker.")
    page_number = f'<w:pgNumType w:start="{start}" w:fmt="{fmt}"/>'.encode()
    members["word/document.xml"] = document_xml.replace(
        marker,
        page_number + marker,
        1,
    )

    with tempfile.NamedTemporaryFile(
        dir=path.parent,
        prefix=".docmergeforge-word-fixture-",
        suffix=".docx",
        delete=False,
    ) as handle:
        temporary = Path(handle.name)
    try:
        with ZipFile(temporary, "w", compression=ZIP_DEFLATED) as output:
            for name, payload in members.items():
                output.writestr(name, payload)
        temporary.replace(path)
    finally:
        temporary.unlink(missing_ok=True)


def _write_fixture(
    path: Path,
    title: str,
    *,
    landscape: bool,
    page_number_start: int,
    page_number_format: str,
) -> None:
    document = Document()
    document.core_properties.title = title
    document.add_heading(title, level=1)
    document.add_paragraph(f"Representative body content for {title}.")
    document.add_paragraph("First numbered item", style="List Number")
    document.add_paragraph("Second numbered item", style="List Number")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Feature"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Source"
    table.cell(1, 1).text = title

    section = document.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Inches(0.70)
        section.right_margin = Inches(0.80)
        section.top_margin = Inches(0.60)
        section.bottom_margin = Inches(0.65)
        section.header_distance = Inches(0.25)
        section.footer_distance = Inches(0.30)
    section.header.paragraphs[0].text = f"Header — {title}"
    section.footer.paragraphs[0].text = f"Footer — {title}"
    document.save(path)
    _inject_page_number_properties(
        path,
        start=page_number_start,
        fmt=page_number_format,
    )


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Run a synthetic Microsoft Word native multi-document merge smoke."
    )
    parser.add_argument("--output-dir", type=Path, default=Path("word-merge-evidence"))
    parser.add_argument("--timeout", type=int, default=900)
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    if args.timeout < 1:
        raise SystemExit("--timeout must be at least one second")

    output_dir: Path = args.output_dir
    output_dir.mkdir(parents=True, exist_ok=True)
    first = output_dir / "word-merge-source-01.docx"
    second = output_dir / "word-merge-source-02.docx"
    merged = output_dir / "word-native-merged.docx"
    evidence_path = output_dir / "word-native-merge-evidence.json"
    for path in (first, second, merged, evidence_path):
        if path.exists():
            raise SystemExit(f"Refusing to overwrite existing Word smoke artifact: {path}")

    _write_fixture(
        first,
        "Word Native Merge Source 1",
        landscape=False,
        page_number_start=1,
        page_number_format="decimal",
    )
    _write_fixture(
        second,
        "Word Native Merge Source 2",
        landscape=True,
        page_number_start=7,
        page_number_format="upperRoman",
    )
    evidence = run_word_merge_acceptance(
        [first, second],
        merged,
        timeout_seconds=args.timeout,
        start_each_on_new_page=True,
    )
    evidence_path.write_text(
        json.dumps(evidence.to_dict(), indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    print(evidence_path)
    return 0 if evidence.accepted else 2


if __name__ == "__main__":
    raise SystemExit(main())
