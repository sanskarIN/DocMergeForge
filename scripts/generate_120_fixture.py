from __future__ import annotations

import argparse
import hashlib
import zipfile
from pathlib import Path


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    args.output.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document
        from pypdf import PdfWriter
    except ImportError as exc:
        raise SystemExit(f"Install project dependencies first: {exc}")

    for part in range(1, 121):
        pdf = args.output / f"SQL_Full_Mastery_Part_{part:03d}_Ram_Sandesh.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        writer.add_metadata({"/Title": f"SQL Full Mastery Part {part}"})
        with pdf.open("wb") as handle:
            writer.write(handle)

        docx_path = args.output / f"SQL_Full_Mastery_Part_{part:03d}_Ram_Sandesh.docx"
        document = Document()
        document.add_heading(f"Part {part}", 1)
        document.add_paragraph(f"Unique regression marker for Part {part}.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Part"
        table.cell(0, 1).text = str(part)
        table.cell(1, 0).text = "Marker"
        table.cell(1, 1).text = hashlib.sha256(str(part).encode()).hexdigest()[:12]
        document.add_paragraph("https://www.github.com/sanskarIN")
        section = document.sections[0]
        section.header.paragraphs[0].text = f"SQL Full Mastery — Part {part}"
        section.footer.paragraphs[0].text = "Made by the Sanskar"
        document.save(docx_path)

        archive = args.output / f"SQL_Full_Mastery_Part_{part:03d}_Companion_Code.zip"
        with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("README.txt", f"Independent companion project for Part {part}.\\n")
            zf.writestr("example.sql", f"-- Part {part}\\nSELECT {part};\\n")

    print(f"Generated 120 PDF, 120 DOCX, and 120 independent companion ZIP fixtures in {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
