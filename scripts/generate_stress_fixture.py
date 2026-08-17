from __future__ import annotations

import argparse
import hashlib
import zipfile
from pathlib import Path


def _positive_int(value: str) -> int:
    number = int(value)
    if number < 1:
        raise argparse.ArgumentTypeError("value must be a positive integer")
    return number


def _deterministic_text(key: str, size: int) -> str:
    chunks: list[str] = []
    length = 0
    counter = 0
    while length < size:
        digest = hashlib.sha256(f"{key}:{counter}".encode()).hexdigest()
        chunks.append(digest)
        length += len(digest)
        counter += 1
    return "".join(chunks)[:size]


def _write_pdf(path: Path, part: int, pages: int, lines_per_page: int) -> None:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    document = canvas.Canvas(str(path), pagesize=LETTER, pageCompression=1)
    width, height = LETTER
    for page_number in range(1, pages + 1):
        document.setFont("Helvetica", 7)
        document.drawString(36, height - 32, f"Stress Part {part} — Page {page_number}")
        for line_number in range(lines_per_page):
            marker = hashlib.sha256(f"pdf:{part}:{page_number}:{line_number}".encode()).hexdigest()
            y = height - 50 - line_number * 12
            document.drawString(36, y, marker)
        document.drawRightString(width - 36, 24, f"Part {part} / Page {page_number}")
        document.showPage()
    document.save()


def _write_docx(path: Path, part: int, paragraphs: int, paragraph_bytes: int) -> None:
    from docx import Document

    document = Document()
    document.add_heading(f"Stress Fixture Part {part}", level=1)
    for paragraph_number in range(1, paragraphs + 1):
        payload = _deterministic_text(
            f"docx:{part}:{paragraph_number}",
            paragraph_bytes,
        )
        document.add_paragraph(f"Marker {part}:{paragraph_number} {payload}")
    section = document.sections[0]
    section.header.paragraphs[0].text = f"DocMergeForge Stress Fixture — Part {part}"
    section.footer.paragraphs[0].text = "Synthetic acceptance fixture"
    document.save(path)


def _write_companion(path: Path, part: int) -> None:
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("README.txt", f"Synthetic companion package for Part {part}.\n")
        archive.writestr("example.sql", f"-- Stress fixture Part {part}\nSELECT {part};\n")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Generate scalable valid PDF/DOCX/companion fixtures for manual stress runs."
    )
    parser.add_argument("output", type=Path)
    parser.add_argument("--parts", type=_positive_int, default=120)
    parser.add_argument("--pdf-pages", type=_positive_int, default=5)
    parser.add_argument("--pdf-lines-per-page", type=_positive_int, default=40)
    parser.add_argument("--docx-paragraphs", type=_positive_int, default=50)
    parser.add_argument("--paragraph-kib", type=_positive_int, default=1)
    args = parser.parse_args(argv)

    if args.pdf_lines_per_page > 55:
        parser.error("--pdf-lines-per-page must be 55 or less for the synthetic page layout")
    if args.paragraph_kib > 1024:
        parser.error("--paragraph-kib must be 1024 or less")

    args.output.mkdir(parents=True, exist_ok=True)
    paragraph_bytes = args.paragraph_kib * 1024
    for part in range(1, args.parts + 1):
        prefix = f"SQL_Full_Mastery_Part_{part:03d}"
        _write_pdf(
            args.output / f"{prefix}_Ram_Sandesh.pdf",
            part,
            args.pdf_pages,
            args.pdf_lines_per_page,
        )
        _write_docx(
            args.output / f"{prefix}_Ram_Sandesh.docx",
            part,
            args.docx_paragraphs,
            paragraph_bytes,
        )
        _write_companion(
            args.output / f"{prefix}_Companion_Code.zip",
            part,
        )

    total_bytes = sum(path.stat().st_size for path in args.output.rglob("*") if path.is_file())
    print(
        "Generated stress fixture: "
        f"parts={args.parts}, pdf_pages_per_part={args.pdf_pages}, "
        f"docx_paragraphs_per_part={args.docx_paragraphs}, "
        f"paragraph_kib={args.paragraph_kib}, source_bytes={total_bytes}"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
