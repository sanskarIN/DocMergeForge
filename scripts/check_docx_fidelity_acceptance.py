from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document

from docmergeforge.docx.fidelity_acceptance import run_fidelity_roundtrip_acceptance


def build_fixture(path: Path) -> None:
    document = Document()
    document.core_properties.title = "DocMergeForge Fidelity Acceptance"
    document.add_heading("Fidelity Acceptance", level=1)
    document.add_paragraph(
        "Representative smoke content with bold, italic, and list formatting."
    )
    formatted = document.add_paragraph()
    formatted.add_run("Bold text").bold = True
    formatted.add_run(" and ")
    formatted.add_run("italic text").italic = True
    document.add_paragraph("First list item", style="List Bullet")
    document.add_paragraph("Second list item", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Feature"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Round-trip"
    table.cell(1, 1).text = "Measured"
    section = document.sections[0]
    section.header.paragraphs[0].text = "DocMergeForge"
    section.footer.paragraphs[0].text = "Fidelity acceptance fixture"
    document.save(path)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Run explicit DOCX fidelity acceptance smoke.")
    parser.add_argument("--mode", choices=("libreoffice", "word"), default="libreoffice")
    parser.add_argument("--output-dir", type=Path, default=Path("fidelity-evidence"))
    parser.add_argument("--timeout", type=int, default=300)
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    if args.timeout < 1:
        raise SystemExit("--timeout must be at least one second")

    output_dir: Path = args.output_dir
    output_dir.mkdir(parents=True, exist_ok=True)
    source = output_dir / "fidelity-source.docx"
    output = output_dir / f"fidelity-{args.mode}-roundtrip.docx"
    evidence_path = output_dir / f"fidelity-{args.mode}-evidence.json"
    for path in (source, output, evidence_path):
        if path.exists():
            raise SystemExit(f"Refusing to overwrite existing acceptance artifact: {path}")

    build_fixture(source)
    evidence = run_fidelity_roundtrip_acceptance(
        source,
        output,
        args.mode,
        timeout_seconds=args.timeout,
    )
    evidence_path.write_text(
        json.dumps(evidence.to_dict(), indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    print(evidence_path)
    return 0 if evidence.accepted else 2


if __name__ == "__main__":
    raise SystemExit(main())
