from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

PACKAGED_SMOKE_ARGUMENT = "--packaged-smoke"


def _run_publication_smoke() -> None:
    from docx import Document
    from pypdf import PdfWriter

    from docmergeforge.app.service import MergeApplicationService
    from docmergeforge.core.models import (
        DocxSettings,
        MergeProject,
        MergeSettings,
        PdfSettings,
    )

    with tempfile.TemporaryDirectory(prefix="docmergeforge-packaged-smoke-") as raw_root:
        root = Path(raw_root)
        source = root / "source"
        output = root / "output"
        source.mkdir()

        pdf_path = source / "Part 1.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        with pdf_path.open("wb") as handle:
            writer.write(handle)

        docx_path = source / "Part 1.docx"
        document = Document()
        document.add_heading("Packaged Smoke Part 1", level=1)
        document.add_paragraph("DocMergeForge packaged publication smoke test.")
        document.save(str(docx_path))

        project = MergeProject(
            name="Packaged Smoke",
            source_folders=[source],
            output_folder=output,
            settings=MergeSettings(
                expected_start=1,
                expected_end=1,
                pdf=PdfSettings(
                    include_title_page=True,
                    page_numbers=True,
                    title="Packaged Smoke",
                ),
                docx=DocxSettings(),
            ),
        )
        artifacts = MergeApplicationService().run_project(project)

        if len(artifacts) != 2:
            raise RuntimeError(
                f"Packaged smoke expected 2 manuscript outputs, got {len(artifacts)}."
            )
        if not all(
            artifact.path.is_file() and artifact.validation_passed for artifact in artifacts
        ):
            raise RuntimeError(
                "Packaged smoke manuscript validation did not complete successfully."
            )
        if not list(output.glob("*_Merge_Manifest.json")):
            raise RuntimeError("Packaged smoke did not create a merge manifest.")
        if not list(output.glob("*_SHA256SUMS.txt")):
            raise RuntimeError("Packaged smoke did not create checksums.")


def _run_packaged_smoke() -> int:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

    from PySide6.QtWidgets import QApplication

    from docmergeforge.diagnostics.logging import configure_logging
    from docmergeforge.settings.config import AppSettings
    from docmergeforge.ui.desktop_entry import ProjectSyncMainWindow
    from docmergeforge.ui.paths import log_path, settings_path
    from docmergeforge.ui.theme import apply_text_scale, apply_theme

    app = QApplication.instance()
    if not isinstance(app, QApplication):
        app = QApplication(sys.argv)
    app.setApplicationName("DocMergeForge")
    app.setOrganizationName("Sanskar")

    settings = AppSettings.load(settings_path())
    configure_logging(log_path(), settings.logging_level)
    apply_theme(app, settings.theme)
    apply_text_scale(app, settings.text_scale_percent)

    window = ProjectSyncMainWindow()
    window.close()
    app.processEvents()
    _run_publication_smoke()
    return 0


def main() -> int:
    if PACKAGED_SMOKE_ARGUMENT in sys.argv[1:]:
        return _run_packaged_smoke()

    from docmergeforge.ui.desktop_entry import main as desktop_main

    return desktop_main()


if __name__ == "__main__":
    raise SystemExit(main())
