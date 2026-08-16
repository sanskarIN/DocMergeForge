from __future__ import annotations

from dataclasses import asdict, dataclass
from pathlib import Path

from docmergeforge.core.models import InputDocument


@dataclass(slots=True, frozen=True)
class PdfComparison:
    source_pages: int
    output_pages: int
    page_count_matches: bool
    part_page_ranges: dict[int, tuple[int, int]]


@dataclass(slots=True, frozen=True)
class DocxCounts:
    paragraphs: int
    tables: int
    inline_shapes: int
    sections: int
    headings: int


@dataclass(slots=True, frozen=True)
class DocxComparison:
    sources: DocxCounts
    output: DocxCounts

    def to_dict(self) -> dict[str, object]:
        return {"sources": asdict(self.sources), "output": asdict(self.output)}


def compare_pdf(inputs: list[InputDocument], output: Path) -> PdfComparison:
    from pypdf import PdfReader

    output_pages = len(PdfReader(str(output), strict=False).pages)
    source_pages = 0
    page_ranges: dict[int, tuple[int, int]] = {}
    cursor = 1
    for item in sorted(inputs, key=lambda value: value.part.number or 10**12):
        pages = item.page_count
        if pages is None:
            pages = len(PdfReader(str(item.path), strict=False).pages)
        start = cursor
        end = cursor + pages - 1
        if item.part.number is not None:
            page_ranges[item.part.number] = (start, end)
        cursor = end + 1
        source_pages += pages
    return PdfComparison(source_pages, output_pages, source_pages == output_pages, page_ranges)


def _docx_counts(path: Path) -> DocxCounts:
    from docx import Document

    document = Document(str(path))
    headings = sum(1 for p in document.paragraphs if p.style and p.style.name.startswith("Heading"))
    return DocxCounts(
        paragraphs=len(document.paragraphs),
        tables=len(document.tables),
        inline_shapes=len(document.inline_shapes),
        sections=len(document.sections),
        headings=headings,
    )


def compare_docx(inputs: list[InputDocument], output: Path) -> DocxComparison:
    source_counts = [_docx_counts(item.path) for item in inputs]
    total = DocxCounts(
        paragraphs=sum(item.paragraphs for item in source_counts),
        tables=sum(item.tables for item in source_counts),
        inline_shapes=sum(item.inline_shapes for item in source_counts),
        sections=sum(item.sections for item in source_counts),
        headings=sum(item.headings for item in source_counts),
    )
    return DocxComparison(total, _docx_counts(output))
